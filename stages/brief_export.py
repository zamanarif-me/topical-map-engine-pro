"""
Stage 9.3: Brief Exporters — CSV and DOCX

Generates downloadable CSV (flattened, one row per brief) and DOCX
(formatted document with all briefs) from a dict of ContentBrief objects.
"""

from __future__ import annotations

import csv
import io
from typing import Iterable

from stages.brief import ContentBrief


# ── CSV ───────────────────────────────────────────────────────────────────────

CSV_COLUMNS = [
    "page_id",
    "page_title",
    "page_type",
    "parent_pillar",
    "central_entity",
    "section_type",
    "information_gain_angle",
    "journey_stage",
    "primary_query",
    "secondary_queries",
    "question_queries",
    "core_attributes",
    "performance_states",
    "failure_modes",
    "dependencies",
    "optimization_levers",
    "heading_outline",
    "must_include",
    "should_include",
    "semantic_variants",
    "recommended_word_count",
    "content_format",
    "reading_level",
    "pov",
    "e_e_a_t_signals",
    "featured_snippet",
    "featured_snippet_section",
    "featured_snippet_format",
    "people_also_ask",
    "schema_markup",
    "semantic_bridges",
    "next_page_id",
    "next_page_title",
    "transition_reason",
    "transition_anchor",
    "author_expertise",
    "experience_signals",
    "trust_signals",
    "ymyl_considerations",
]


def _join(items: Iterable, sep: str = " | ") -> str:
    return sep.join(str(x) for x in items if x)


def _brief_to_row(brief: ContentBrief) -> dict:
    headings = " > ".join(
        f"{h.level}: {h.text}" for h in brief.headings if h.text
    )
    bridges = " || ".join(
        f"[{float(b.relationship_strength):.2f}] {b.link_destination} via {b.shared_entity} :: {b.anchor_suggestion}"
        for b in brief.semantic_bridges
    )
    return {
        "page_id": brief.page_id,
        "page_title": brief.page_title,
        "page_type": brief.page_type,
        "parent_pillar": brief.parent_pillar or "",
        "central_entity": brief.central_entity,
        "section_type": brief.section_type,
        "information_gain_angle": brief.information_gain_angle,
        "journey_stage": brief.queries.journey_stage,
        "primary_query": brief.queries.primary_query,
        "secondary_queries": _join(brief.queries.secondary_queries),
        "question_queries": _join(brief.queries.question_queries),
        "core_attributes": _join(brief.entity_attribute_map.core_attributes),
        "performance_states": _join(brief.entity_attribute_map.performance_states),
        "failure_modes": _join(brief.entity_attribute_map.failure_modes),
        "dependencies": _join(brief.entity_attribute_map.dependencies),
        "optimization_levers": _join(brief.entity_attribute_map.optimization_levers),
        "heading_outline": headings,
        "must_include": _join(brief.nlp_terms.must_include),
        "should_include": _join(brief.nlp_terms.should_include),
        "semantic_variants": _join(brief.nlp_terms.semantic_variants),
        "recommended_word_count": brief.content_specs.recommended_word_count,
        "content_format": brief.content_specs.content_format,
        "reading_level": brief.content_specs.reading_level,
        "pov": brief.content_specs.pov,
        "e_e_a_t_signals": _join(brief.content_specs.e_e_a_t_signals),
        "featured_snippet": "Yes" if brief.serp_target.featured_snippet else "No",
        "featured_snippet_section": brief.serp_target.featured_snippet_section or "",
        "featured_snippet_format": brief.serp_target.featured_snippet_format or "",
        "people_also_ask": _join(brief.serp_target.people_also_ask),
        "schema_markup": brief.serp_target.schema_markup,
        "semantic_bridges": bridges,
        "next_page_id": brief.next_destination.next_page_id,
        "next_page_title": brief.next_destination.next_page_title,
        "transition_reason": brief.next_destination.transition_reason,
        "transition_anchor": brief.next_destination.transition_anchor,
        "author_expertise": brief.eeat_requirements.author_expertise,
        "experience_signals": _join(brief.eeat_requirements.experience_signals),
        "trust_signals": _join(brief.eeat_requirements.trust_signals),
        "ymyl_considerations": brief.eeat_requirements.ymyl_considerations or "",
    }


def briefs_to_csv(briefs: dict[str, ContentBrief]) -> bytes:
    """Render all briefs as a single CSV — one row per brief."""
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=CSV_COLUMNS, extrasaction="ignore")
    writer.writeheader()
    for brief in briefs.values():
        writer.writerow(_brief_to_row(brief))
    return buf.getvalue().encode("utf-8-sig")  # BOM for Excel


# ── DOCX ──────────────────────────────────────────────────────────────────────

def briefs_to_docx(briefs: dict[str, ContentBrief], title: str = "Content Briefs") -> bytes:
    """Render all briefs as a single Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed. Add 'python-docx>=1.1.0' to requirements.txt"
        ) from e

    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title_p = doc.add_heading(title, level=0)

    doc.add_paragraph(f"{len(briefs)} content briefs generated by the Topical Map Engine.")
    doc.add_paragraph()

    # ── One section per brief ─────────────────────────────────────────────────
    for i, (page_id, brief) in enumerate(briefs.items()):
        if i > 0:
            doc.add_page_break()

        doc.add_heading(brief.page_title or page_id, level=1)

        # Meta box
        meta = doc.add_paragraph()
        meta.add_run("Page ID: ").bold = True
        meta.add_run(brief.page_id + "\n")
        meta.add_run("Type: ").bold = True
        meta.add_run(f"{brief.page_type}  |  ")
        meta.add_run("Section: ").bold = True
        meta.add_run(f"{brief.section_type}  |  ")
        meta.add_run("Word Count: ").bold = True
        meta.add_run(f"{brief.content_specs.recommended_word_count:,}\n")
        meta.add_run("Journey Stage: ").bold = True
        meta.add_run(brief.queries.journey_stage)

        # Information Gain
        doc.add_heading("Information Gain Angle", level=2)
        doc.add_paragraph(brief.information_gain_angle or "—")

        # Target Queries
        doc.add_heading("Target Queries", level=2)
        p = doc.add_paragraph()
        p.add_run("Primary: ").bold = True
        p.add_run(brief.queries.primary_query or "—")
        if brief.queries.secondary_queries:
            doc.add_paragraph("Secondary queries:", style="Intense Quote")
            for q in brief.queries.secondary_queries:
                doc.add_paragraph(q, style="List Bullet")
        if brief.queries.question_queries:
            doc.add_paragraph("Questions to answer:", style="Intense Quote")
            for q in brief.queries.question_queries:
                doc.add_paragraph(q, style="List Bullet")

        # Entity Attribute Map
        doc.add_heading("Entity Attribute Map", level=2)
        eam = brief.entity_attribute_map
        _docx_kv_list(doc, [
            ("Core attributes", eam.core_attributes),
            ("Performance states", eam.performance_states),
            ("Failure modes", eam.failure_modes),
            ("Dependencies", eam.dependencies),
            ("Optimization levers", eam.optimization_levers),
        ])

        # Heading Structure
        doc.add_heading("Heading Structure", level=2)
        for h in brief.headings:
            lvl = h.level if h.level else "H2"
            depth = 0
            if len(lvl) > 1 and lvl[1].isdigit():
                depth = max(0, int(lvl[1]) - 1)
            indent = "    " * depth
            p = doc.add_paragraph()
            p.add_run(f"{indent}{lvl}: ").bold = True
            p.add_run(h.text)
            if h.semantic_purpose:
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Pt(18 * (depth + 1))
                run = sp.add_run(h.semantic_purpose)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x8A)

        # NLP Terms
        doc.add_heading("NLP Terms", level=2)
        _docx_kv_list(doc, [
            ("Must include (3-5x)", brief.nlp_terms.must_include),
            ("Should include (1-2x)", brief.nlp_terms.should_include),
            ("Semantic variants", brief.nlp_terms.semantic_variants),
        ])

        # Content Specs
        doc.add_heading("Content Specifications", level=2)
        cs = brief.content_specs
        _docx_kv_list(doc, [
            ("Format", [cs.content_format]),
            ("Reading level", [cs.reading_level]),
            ("POV", [cs.pov]),
            ("E-E-A-T signals", cs.e_e_a_t_signals),
        ])

        # SERP target
        doc.add_heading("SERP Target", level=2)
        st = brief.serp_target
        _docx_kv_list(doc, [
            ("Featured snippet", ["Yes" if st.featured_snippet else "No"]),
            ("Snippet section", [st.featured_snippet_section or "—"]),
            ("Snippet format", [st.featured_snippet_format or "—"]),
            ("Schema markup", [st.schema_markup]),
            ("People Also Ask", st.people_also_ask),
        ])

        # Semantic bridges
        doc.add_heading("Semantic Bridges", level=2)
        if not brief.semantic_bridges:
            doc.add_paragraph("—")
        else:
            for b in brief.semantic_bridges:
                strength = float(b.relationship_strength) if b.relationship_strength else 0.0
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{strength:.2f}] ").bold = True
                p.add_run(f"→ {b.link_destination}")
                p.add_run(f" via {b.shared_entity}").italic = True
                doc.add_paragraph(f"    Bridge point: {b.bridge_point}")
                doc.add_paragraph(f"    Anchor: {b.anchor_suggestion}")

        # Next destination
        doc.add_heading("Next Destination", level=2)
        nd = brief.next_destination
        p = doc.add_paragraph()
        p.add_run("Next page: ").bold = True
        p.add_run(f"{nd.next_page_title}  ({nd.next_page_id})\n")
        p.add_run("Why: ").bold = True
        p.add_run(f"{nd.transition_reason}\n")
        p.add_run("CTA: ").bold = True
        p.add_run(nd.transition_anchor).italic = True

        # E-E-A-T
        doc.add_heading("E-E-A-T Requirements", level=2)
        ee = brief.eeat_requirements
        p = doc.add_paragraph()
        p.add_run("Author expertise: ").bold = True
        p.add_run(ee.author_expertise or "—")
        _docx_kv_list(doc, [
            ("Experience signals", ee.experience_signals),
            ("Trust signals", ee.trust_signals),
        ])
        if ee.ymyl_considerations:
            p = doc.add_paragraph()
            p.add_run("YMYL considerations: ").bold = True
            p.add_run(ee.ymyl_considerations)

        # Quality checklist
        if brief.quality_checklist:
            doc.add_heading("Quality Checklist", level=2)
            for key, val in brief.quality_checklist.items():
                mark = "☑" if val else "☐"
                doc.add_paragraph(f"{mark} {key}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_kv_list(doc, pairs: list[tuple[str, list]]) -> None:
    """Render a key→list block as bold-key + bulleted values."""
    for label, items in pairs:
        items = [x for x in (items or []) if x]
        if not items:
            continue
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(", ".join(str(x) for x in items))
